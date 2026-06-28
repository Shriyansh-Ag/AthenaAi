import docx
from typing import BinaryIO
from app.models.document import ExtractedDocument, DocumentMetadata, ExtractedPage, ContentBlock
from app.services.extractors.base_extractor import BaseExtractor
from app.utils.language_utils import detect_language

class DocxExtractor(BaseExtractor):
    def extract(self, file_stream: BinaryIO, filename: str) -> ExtractedDocument:
        doc = docx.Document(file_stream)
        
        metadata = DocumentMetadata(
            title=doc.core_properties.title or filename,
            author=doc.core_properties.author or "",
            creation_date=str(doc.core_properties.created) if doc.core_properties.created else "",
            file_type="docx"
        )
        
        content_blocks = []
        headings = []
        all_text = ""
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            all_text += text + " "
            
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split(" ")[1])
                except (IndexError, ValueError):
                    level = 1
                    
                cb = ContentBlock(type="heading", text=text, level=level)
                headings.append(cb)
            else:
                cb = ContentBlock(type="paragraph", text=text)
                
            content_blocks.append(cb)
            
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            cb = ContentBlock(type="table", metadata={"data": table_data})
            content_blocks.append(cb)
            
        # DOCX doesn't have a strict concept of pages easily accessible via python-docx
        page = ExtractedPage(page_number=1, content=content_blocks)
        
        metadata.language = detect_language(all_text)
        
        return ExtractedDocument(
            metadata=metadata,
            pages=[page],
            headings=headings
        )
